import streamlit as st
import logging
import asyncio
import time
import os
import unicodedata
import re
import spacy
import difflib
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Bibliotecas para OCR e imagem
from pdf2image import convert_from_path
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter

# Criptografia de dados (exemplo simples)
from cryptography.fernet import Fernet

# Configuração básica de logs
logging.basicConfig(level=logging.ERROR)

# Configurar o caminho do Tesseract (ajuste conforme necessário)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'  # Para Linux

# Ajuste para Windows no loop de eventos assíncronos
if os.name == 'nt':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

LOGIN_URL = "https://sei.anvisa.gov.br/sip/login.php?sigla_orgao_sistema=ANVISA&sigla_sistema=SEI"

###############################################################################
# Criptografia básica (chave em memória)
###############################################################################
secret_key = Fernet.generate_key()
cipher_suite = Fernet(secret_key)

###############################################################################
# Funções de Validação de CPF e CNPJ
###############################################################################
def validar_cpf(cpf: str) -> bool:
    cpf = re.sub(r"\D", "", cpf)
    if len(cpf) != 11:
        return False
    if cpf in [str(i)*11 for i in range(10)]:
        return False
    for i in range(9, 11):
        soma = 0
        for j in range(0, i):
            soma += int(cpf[j]) * ((i+1) - j)
        resto = (soma * 10) % 11
        resto = 0 if resto == 10 else resto
        if resto != int(cpf[i]):
            return False
    return True

def validar_cnpj(cnpj: str) -> bool:
    cnpj = re.sub(r"\D", "", cnpj)
    if len(cnpj) != 14:
        return False
    if cnpj in [str(i)*14 for i in range(10)]:
        return False
    
    def calc_dv(cnpj_parcial):
        peso = [6,7,8,9,2,3,4,5]
        soma = 0
        for i, digit in enumerate(cnpj_parcial[::-1]):
            soma += int(digit) * peso[i % len(peso)]
        resto = soma % 11
        return '0' if resto < 2 else str(11 - resto)
    
    dv1 = calc_dv(cnpj[:-2])
    dv2 = calc_dv(cnpj[:-2] + dv1)
    return (dv1 == cnpj[-2]) and (dv2 == cnpj[-1])

###############################################################################
# Funções relacionadas ao Playwright
###############################################################################
def create_browser_context(headless=True):
    download_dir = os.path.join(os.getcwd(), "downloads")
    os.makedirs(download_dir, exist_ok=True)
    
    user_data_dir = os.path.join(os.getcwd(), "user_data")
    os.makedirs(user_data_dir, exist_ok=True)
    
    playwright = sync_playwright().start()
    
    context = playwright.chromium.launch_persistent_context(
        user_data_dir=user_data_dir,
        headless=headless,
        accept_downloads=True,
        downloads_path=download_dir
    )
    page = context.new_page()
    return playwright, context, page

def wait_for_element(page, selector, timeout=20000):
    try:
        element = page.wait_for_selector(selector, timeout=timeout)
        if element:
            return element
    except PlaywrightTimeoutError:
        logging.error(f"Elemento {selector} não encontrado na página.")
        raise Exception(f"Elemento {selector} não encontrado na página.")
    return None

def handle_download(download, download_dir):
    os.makedirs(download_dir, exist_ok=True)
    download_path = os.path.join(download_dir, download.suggested_filename)
    download.save_as(download_path)
    logging.info(f"Download salvo em: {download_path}")
    return download_path

def handle_alert(page):
    try:
        dialog = page.expect_event("dialog", timeout=5000)
        if dialog:
            alert_text = dialog.message
            logging.warning(f"Alerta inesperado encontrado: {alert_text}")
            dialog.accept()
            return alert_text
    except PlaywrightTimeoutError:
        return None

def login(page, username_encrypted, password_encrypted):
    username = cipher_suite.decrypt(username_encrypted).decode('utf-8')
    password = cipher_suite.decrypt(password_encrypted).decode('utf-8')
    
    page.goto(LOGIN_URL)
    
    user_field = wait_for_element(page, "#txtUsuario")
    if user_field:
        user_field.fill(username)
    else:
        raise Exception("Campo de usuário não encontrado.")
    
    password_field = wait_for_element(page, "#pwdSenha")
    if password_field:
        password_field.fill(password)
    else:
        raise Exception("Campo de senha não encontrado.")
    
    login_button = wait_for_element(page, "#sbmAcessar")
    if login_button:
        login_button.click()
    else:
        raise Exception("Botão de login não encontrado.")
    
    try:
        page.wait_for_load_state("networkidle", timeout=20000)
    except PlaywrightTimeoutError:
        raise Exception("Login pode não ter sido realizado com sucesso.")

def access_process(page, process_number):
    try:
        search_field = wait_for_element(page, "#txtPesquisaRapida", timeout=40000)
        search_field.fill(process_number)
        search_field.press("Enter")
        time.sleep(5)
    except Exception as e:
        raise Exception(f"Erro ao acessar o processo: {e}")

IFRAME_VISUALIZACAO_ID = "ifrVisualizacao"
BUTTON_XPATH_GERAR_PDF = '//*[@id="divArvoreAcoes"]/a[7]/img'
BUTTON_XPATH_DOWNLOAD_OPTION = '//*[@id="divInfraBarraComandosSuperior"]/button[1]'

def generate_and_download_pdf(page, download_dir):
    try:
        iframe_element = page.wait_for_selector(f'iframe#{IFRAME_VISUALIZACAO_ID}', timeout=10000)
        if not iframe_element:
            raise Exception(f"Iframe com ID {IFRAME_VISUALIZACAO_ID} não encontrado.")
        
        iframe = iframe_element.content_frame()
        if not iframe:
            raise Exception("Não foi possível acessar o conteúdo do iframe.")
        
        gerar_pdf_button = iframe.wait_for_selector(f'xpath={BUTTON_XPATH_GERAR_PDF}', timeout=10000)
        if not gerar_pdf_button:
            raise Exception("Botão para gerar PDF não encontrado.")
        gerar_pdf_button.click()
        time.sleep(2)
        
        download_option_button = iframe.wait_for_selector(f'xpath={BUTTON_XPATH_DOWNLOAD_OPTION}', timeout=10000)
        if not download_option_button:
            raise Exception("Botão de opção de download não encontrado.")
        
        with page.expect_download(timeout=60000) as download_info_option:
            download_option_button.click()
        download_option = download_info_option.value
        download_option_path = handle_download(download_option, download_dir)
        
        return download_option_path
    
    except PlaywrightTimeoutError:
        raise Exception("Timeout ao gerar o PDF do processo.")
    except Exception as e:
        raise Exception(f"Erro ao gerar o PDF do processo: {e}")
    finally:
        time.sleep(5)

def process_notification(username_encrypted, password_encrypted, process_number, headless=True):
    download_dir = os.path.join(os.getcwd(), "downloads")
    playwright, context, page = create_browser_context(headless=headless)
    
    try:
        login(page, username_encrypted, password_encrypted)
        access_process(page, process_number)
        download_path = generate_and_download_pdf(page, download_dir)
        return download_path
    except Exception as e:
        logging.error(f"Erro durante o processamento: {e}")
        raise e
    finally:
        context.close()
        playwright.stop()

###############################################################################
# Extração de texto e OCR (atualizado)
###############################################################################
def normalize_text(text):
    if not isinstance(text, str):
        return text
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

def corrigir_texto(texto):
    substituicoes = {
        'Ã©': 'é',
        'Ã§Ã£o': 'ção',
        'Ã³': 'ó',
        'Ã': 'à',
        'â€“': '–',
        'â€”': '—',
        'Ãº': 'ú',
        'Ãª': 'ê',
        'Ã£o': 'ão',
        'â€œ': '"',
        'â€': '"',
        'Ã¡': 'á',
        'Ã¢': 'â',
        'Ã­': 'í',
        'Ã´': 'ô',
        'Ã§': 'ç',
    }
    for errado, correto in substituicoes.items():
        texto = texto.replace(errado, correto)
    return texto

def extract_text_with_pypdf2(pdf_path):
    """
    Primeiro tenta extrair texto via PyPDF2, sem OCR.
    Se der certo, retorna o texto.
    Caso não encontre nada, retorna string vazia.
    """
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        
        if text.strip():
            text = corrigir_texto(normalize_text(text))
            return text.strip()
        else:
            return ''
    except:
        return ''

def extract_text_with_context(image_path, file_origin, lang='por'):
    """
    Extrai texto de uma imagem com Tesseract e localiza endereços básicos via regex.
    - Filtra endereços com menos de 15 caracteres (campo 'endereco').
    - Adiciona 'file_origin' em cada endereço apenas como referência/visão do usuário.
    """
    try:
        custom_config = f"--psm 6 --oem 3 -l {lang}"
        image = Image.open(image_path)
        text_page = pytesseract.image_to_string(image, config=custom_config)

        text_page = corrigir_texto(normalize_text(text_page))

        # Regex simples para capturar Endereco, Cidade, Bairro, Estado e CEP
        endereco_pattern = r"Endere[c|ç]o[:\s]+([\w\s.,/\-ºª]+)"
        cidade_pattern   = r"Cidade[:\s]+([\w\s]+)"
        bairro_pattern   = r"Bairro[:\s]+([\w\s]+)"
        estado_pattern   = r"Estado[:\s]+([A-Z]{2})"
        cep_pattern      = r"CEP[:\s]+([\d.\-]+)"

        enderecos_encontrados = []

        end_matches = re.findall(endereco_pattern, text_page, flags=re.IGNORECASE)
        cid_matches = re.findall(cidade_pattern, text_page, flags=re.IGNORECASE)
        bai_matches = re.findall(bairro_pattern, text_page, flags=re.IGNORECASE)
        uf_matches  = re.findall(estado_pattern, text_page, flags=re.IGNORECASE)
        cep_matches = re.findall(cep_pattern, text_page, flags=re.IGNORECASE)

        max_len = max(len(end_matches), len(cid_matches), len(bai_matches), len(uf_matches), len(cep_matches))

        for i in range(max_len):
            endereco_val = end_matches[i] if i < len(end_matches) else "[Não informado]"
            cidade_val   = cid_matches[i] if i < len(cid_matches) else "[Não informado]"
            bairro_val   = bai_matches[i] if i < len(bai_matches) else "[Não informado]"
            estado_val   = uf_matches[i]  if i < len(uf_matches)  else "[Não informado]"
            cep_val      = cep_matches[i] if i < len(cep_matches) else "[Não informado]"

            # Excluir endereços com menos de 15 caracteres
            if len(endereco_val.strip()) < 15:
                continue

            enderecos_encontrados.append({
                "endereco": endereco_val,
                "cidade": cidade_val,
                "bairro": bairro_val,
                "estado": estado_val,
                "cep": cep_val,
                "source": file_origin  # apenas exibição em tela
            })

        return text_page, enderecos_encontrados

    except Exception as e:
        logging.error(f"Erro ao processar a imagem {image_path}: {e}")
        return "", []

def ocr_extract(pdf_path, psm_mode=6, oem_mode=3):
    """
    Extrai texto via OCR de cada página do PDF (convertida em imagem).
    Retorna todo o texto concatenado e também uma lista de endereços
    encontrados por regex, com respectivo 'source'.
    """
    text_total = ""
    enderecos_totais = []

    try:
        pages = convert_from_path(pdf_path, dpi=300, fmt='jpeg')

        for idx, page in enumerate(pages, start=1):
            gray = page.convert('L')
            enhancer = ImageEnhance.Contrast(gray)
            gray = enhancer.enhance(2.0)
            threshold = gray.point(lambda x: 0 if x < 128 else 255, '1')
            threshold = threshold.filter(ImageFilter.MedianFilter())

            temp_filename = f"temp_page_{idx}.jpg"
            threshold.save(temp_filename, "JPEG")

            file_origin = f"{os.path.basename(pdf_path)} - Página {idx}"
            text_page, enderecos_page = extract_text_with_context(temp_filename, file_origin, lang='por')

            text_total += text_page + "\n"
            enderecos_totais.extend(enderecos_page)

            if os.path.exists(temp_filename):
                os.remove(temp_filename)

    except Exception as e:
        st.error(f"Erro durante o OCR: {e}")

    text_total = corrigir_texto(normalize_text(text_total))
    return text_total, enderecos_totais

def extract_text_with_best_ocr(pdf_path):
    """
    Tenta extrair texto sem OCR (PyPDF2).
    Se não conseguir, faz OCR em cada página.
    Retorna o texto final e a lista de endereços extraídos (com .source).
    """
    extracted_text = extract_text_with_pypdf2(pdf_path)
    if extracted_text.strip():
        # Se extraiu com PyPDF2, não faz OCR
        return extracted_text, []
    
    # Caso contrário, faz OCR
    text_ocr, enderecos_ocr = ocr_extract(pdf_path, psm_mode=6, oem_mode=3)
    if len(text_ocr) > 0:
        return text_ocr, enderecos_ocr

    return "", []

###############################################################################
# Formatação e extração de dados
###############################################################################
def format_cnpj(cnpj):
    digits = re.sub(r'\D', '', cnpj)
    if len(digits) != 14:
        return cnpj
    return f"{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:14]}"

def format_cpf(cpf):
    digits = re.sub(r'\D', '', cpf)
    if len(digits) != 11:
        return cpf
    return f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:11]}"

def extract_process_number(file_name):
    base_name = os.path.splitext(file_name)[0]
    if base_name.startswith("SEI"):
        base_name = base_name[3:].strip()
    
    digits = re.sub(r'\D', '', base_name)
    if len(digits) != 15:
        return base_name
    return f"{digits[:5]}.{digits[5:11]}/{digits[11:15]}-{digits[14:]}"

def extract_information_spacy(text):
    """
    Exemplo de extração com spacy (nomes, e-mails, etc.).
    """
    doc = nlp(text)
    info = {
        "nome_autuado": None,
        "cpf": None,
        "cnpj": None,
        "socios_advogados": [],
        "emails": [],
    }
    
    for ent in doc.ents:
        if ent.label_ in ["PER", "ORG"]:
            if not info["nome_autuado"]:
                info["nome_autuado"] = ent.text.strip()
        elif ent.label_ == "EMAIL":
            info["emails"].append(ent.text.strip())
    
    # Regex para CNPJ/CPF
    cnpj_pattern = r"CNPJ:\s*([\d./-]{14,18})"
    cpf_pattern  = r"CPF:\s*([\d./-]{11,14})"

    cnpj_match = re.search(cnpj_pattern, text)
    cpf_match  = re.search(cpf_pattern, text)
    
    if cnpj_match:
        cnpj_puro = re.sub(r'\D', '', cnpj_match.group(1))
        if validar_cnpj(cnpj_puro):
            info["cnpj"] = format_cnpj(cnpj_match.group(1))
    
    if cpf_match:
        cpf_puro = re.sub(r'\D', '', cpf_match.group(1))
        if validar_cpf(cpf_puro):
            info["cpf"] = format_cpf(cpf_match.group(1))
    
    # Sócios / advogados
    socios_adv_pattern = r"(?:Sócio|Advogado|Responsável|Representante Legal):\s*([\w\s]+)"
    info["socios_advogados"] = re.findall(socios_adv_pattern, text) or []
    
    return info

def extract_addresses_with_source(text):
    """
    Exemplo: extrai endereços com a 'source' baseada em 'AR' ou 'AIS' no texto.
    + Filtrar endereços < 15 caracteres.
    """
    doc = nlp(text)
    page_blocks = text.split("\f")
    
    addresses = []
    
    endereco_pattern = r"(?:Endereço|End|Endereco):\s*([\w\s.,ºª-]+)"
    cidade_pattern   = r"Cidade:\s*([\w\s]+(?: DE [\w\s]+)?)"
    bairro_pattern   = r"Bairro:\s*([\w\s]+)"
    estado_pattern   = r"Estado:\s*([A-Z]{2})"
    cep_pattern      = r"CEP:\s*(\d{2}\.\d{3}-\d{3}|\d{5}-\d{3})"
    
    for block in page_blocks:
        block_clean = block.strip()
        block_source = "Desconhecido"
        if re.search(r"\bAR\b", block_clean, re.IGNORECASE):
            block_source = "AR"
        elif re.search(r"\bAIS\b", block_clean, re.IGNORECASE):
            block_source = "AIS"
        
        endereco_matches = re.findall(endereco_pattern, block_clean, re.IGNORECASE)
        cidade_matches   = re.findall(cidade_pattern, block_clean, re.IGNORECASE)
        bairro_matches   = re.findall(bairro_pattern, block_clean, re.IGNORECASE)
        estado_matches   = re.findall(estado_pattern, block_clean, re.IGNORECASE)
        cep_matches      = re.findall(cep_pattern, block_clean, re.IGNORECASE)
        
        max_len = max(
            len(endereco_matches),
            len(cidade_matches),
            len(bairro_matches),
            len(estado_matches),
            len(cep_matches)
        )
        
        for i in range(max_len):
            end_str = endereco_matches[i].strip() if i < len(endereco_matches) else "[Não informado]"
            cid_str = cidade_matches[i].strip()   if i < len(cidade_matches)   else "[Não informado]"
            bai_str = bairro_matches[i].strip()   if i < len(bairro_matches)   else "[Não informado]"
            uf_str  = estado_matches[i].strip()   if i < len(estado_matches)   else "[Não informado]"
            cep_str = cep_matches[i].strip()      if i < len(cep_matches)      else "[Não informado]"
            
            if len(end_str) < 15:
                continue
            
            addresses.append({
                "endereco": end_str,
                "cidade": cid_str,
                "bairro": bai_str,
                "estado": uf_str,
                "cep": cep_str,
                "source": block_source
            })
    
    return addresses

def normalize_address(address):
    address = unicodedata.normalize('NFKD', address).encode('ASCII', 'ignore').decode('utf-8')
    address = re.sub(r'[^\w\s]', '', address)
    address = re.sub(r'\s+', ' ', address)
    return address.lower().strip()

def extract_all_emails(emails):
    return list(set(emails))

###############################################################################
# Modelos Word
###############################################################################
def adicionar_paragrafo(doc, texto="", negrito=False, tamanho=12):
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run(texto)
    run.bold = negrito
    run.font.size = Pt(tamanho)
    return paragrafo

def _gerar_modelo_1(doc, info, enderecos, numero_processo, email_selecionado):
    """
    Gera o Documento Word no Modelo 1.
    
    :param doc: Objeto Document do python-docx.
    :param info: Dicionário com informações extraídas.
    :param enderecos: Lista de dicionários com endereços.
    :param numero_processo: Número do processo formatado.
    :param email_selecionado: Email selecionado pelo usuário.
    """
    try:
        # Cabeçalho do documento
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        nome_autuado = info.get('nome_autuado', '[Nome não informado]')
        cnpj = info.get('cnpj', '')
        cpf = info.get('cpf', '')
        if cnpj:
            identificador = f"CNPJ: {cnpj}"
        elif cpf:
            identificador = f"CPF: {cpf}"
        else:
            identificador = "CNPJ/CPF: [Não informado]"
        adicionar_paragrafo(doc, f"{nome_autuado} – {identificador}")
        doc.add_paragraph("\n")

        # Informações de endereço
        for endereco in enderecos:
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Cidade: {endereco.get('cidade', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Bairro: {endereco.get('bairro', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Estado: {endereco.get('estado', '[Não informado]')}")
            adicionar_paragrafo(doc, f"CEP: {endereco.get('cep', '[Não informado]')}")
            doc.add_paragraph("\n")

        # Assunto e Referência em negrito
        adicionar_paragrafo(doc, 
            "Assunto: Decisão de 1ª instância proferida pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias.", 
            negrito=True
        )
        adicionar_paragrafo(doc, 
            f"Referência: Processo Administrativo Sancionador nº: {numero_processo} ", 
            negrito=True
        )
        doc.add_paragraph("\n")

        # Corpo do documento
        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, 
            "Informamos que foi proferido julgamento pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias no processo administrativo sancionador em referência, conforme decisão em anexo."
        )
        doc.add_paragraph("\n")

        # Seções do documento com formatação
        adicionar_paragrafo(doc, "O QUE FAZER SE A DECISÃO TIVER APLICADO MULTA?", negrito=True)
        adicionar_paragrafo(doc, 
            "Sendo aplicada a penalidade de multa, esta notificação estará acompanhada de boleto bancário, que deverá ser pago até o vencimento."
        )
        adicionar_paragrafo(doc, 
            "O valor da multa poderá ser pago com 20% de desconto caso seja efetuado em até 20 dias contados de seu recebimento. "
            "Incorrerá em ilegalidade o usufruto do desconto em data posterior ao prazo referido, mesmo que a data impressa no boleto permita pagamento, "
            "sendo a diferença cobrada posteriormente pela Gerência de Gestão de Arrecadação (GEGAR). "
            "O pagamento da multa implica em desistência tácita do recurso, conforme art. 21 da Lei nº 6.437/1977."
        )
        adicionar_paragrafo(doc, 
            "O não pagamento do boleto sem que haja interposição de recurso, acarretará, sucessivamente: "
            "i) a inscrição do devedor no Cadastro Informativo de Crédito não Quitado do Setor Público Federal (CADIN); "
            "ii) a inscrição do débito em dívida ativa da União; iii) o ajuizamento de ação de execução fiscal contra o devedor; "
            "e iv) a comunicação aos cartórios de registros de imóveis, dos devedores inscritos em dívida ativa ou execução fiscal."
        )
        adicionar_paragrafo(doc, 
            "Esclarecemos que o valor da multa foi atualizado pela taxa Selic acumulada nos termos do art. 37-A da Lei 10.522/2002 "
            "e no art. 5º do Decreto-Lei 1.736/79."
        )
        doc.add_paragraph("\n")

        adicionar_paragrafo(doc, "COMO FAÇO PARA INTERPOR RECURSO DA DECISÃO?", negrito=True)
        adicionar_paragrafo(doc, 
            "Havendo interesse na interposição de recurso administrativo, este poderá ser interposto no prazo de 20 dias contados do recebimento desta notificação, "
            "conforme disposto no art. 9º da RDC nº 266/2019."
        )
        adicionar_paragrafo(doc, 
            "O protocolo do recurso deverá ser feito exclusivamente, por meio de peticionamento intercorrente no processo indicado no campo assunto desta notificação, "
            "pelo Sistema Eletrônico de Informações (SEI). Para tanto, é necessário, primeiramente, fazer o cadastro como usuário externo SEI-Anvisa. "
            "Acesse o portal da Anvisa https://www.gov.br/anvisa/pt-br > Sistemas > SEI > Acesso para Usuários Externos (SEI) e siga as orientações. "
            "Para maiores informações, consulte o Manual do Usuário Externo Sei-Anvisa, que está disponível em https://www.gov.br/anvisa/pt-br/sistemas/sei."
        )
        doc.add_paragraph("\n")

        adicionar_paragrafo(doc, "QUAIS DOCUMENTOS DEVEM ACOMPANHAR O RECURSO?", negrito=True)
        adicionar_paragrafo(doc, "a) Autuado pessoa jurídica:")
        adicionar_paragrafo(doc, "1. Contrato ou estatuto social da empresa, com a última alteração;")
        adicionar_paragrafo(doc, 
            "2. Procuração e documento de identificação do outorgado (advogado ou representante), caso constituído para atuar no processo. "
            "Somente serão aceitas procurações e substabelecimentos assinados eletronicamente, com certificação digital no padrão da "
            "Infraestrutura de Chaves Públicas Brasileira (ICP-Brasil) ou pelo assinador Gov.br."
        )
        adicionar_paragrafo(doc, 
            "3. Ata de eleição da atual diretoria quando a procuração estiver assinada por diretor que não conste como sócio da empresa;"
        )
        adicionar_paragrafo(doc, 
            "4. No caso de contestação sobre o porte da empresa considerado para a dosimetria da pena de multa: comprovação do porte econômico "
            "referente ao ano em que foi proferida a decisão (documentos previstos no art. 50 da RDC nº 222/2006)."
        )
        adicionar_paragrafo(doc, "b) Autuado pessoa física:")
        adicionar_paragrafo(doc, "1. Documento de identificação do autuado;")
        adicionar_paragrafo(doc, 
            "2. Procuração e documento de identificação do outorgado (advogado ou representante), caso constituído para atuar no processo."
        )
        adicionar_paragrafo(doc, f"\nInformações de contato: {email_selecionado}")

    except Exception as e:
        st.error(f"Erro ao gerar o documento no modelo 1: {e}")

def _gerar_modelo_2(doc, info, enderecos, numero_processo, motivo_revisao, data_decisao, data_recebimento_notificacao, data_extincao=None, email_selecionado=None):
    """
    Gera o Documento Word no Modelo 2.
    
    :param doc: Objeto Document do python-docx.
    :param info: Dicionário com informações extraídas.
    :param enderecos: Lista de dicionários com endereços.
    :param numero_processo: Número do processo formatado.
    :param motivo_revisao: Motivo da revisão da decisão.
    :param data_decisao: Data da decisão original.
    :param data_recebimento_notificacao: Data de recebimento da notificação.
    :param data_extincao: Data de extinção da empresa (se aplicável).
    :param email_selecionado: Email selecionado pelo usuário.
    """
    try:
        # Cabeçalho do documento
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        nome_autuado = info.get('nome_autuado', '[Nome não informado]')
        cnpj = info.get('cnpj', '')
        cpf = info.get('cpf', '')
        if cnpj:
            identificador = f"CNPJ: {cnpj}"
        elif cpf:
            identificador = f"CPF: {cpf}"
        else:
            identificador = "CNPJ/CPF: [Não informado]"
        adicionar_paragrafo(doc, f"{nome_autuado} – {identificador}")
        doc.add_paragraph("\n")

        # Informações de endereço
        for endereco in enderecos:
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Cidade: {endereco.get('cidade', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Bairro: {endereco.get('bairro', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Estado: {endereco.get('estado', '[Não informado]')}")
            adicionar_paragrafo(doc, f"CEP: {endereco.get('cep', '[Não informado]')}")
            doc.add_paragraph("\n")

        # Assunto e Referência em negrito
        adicionar_paragrafo(doc, 
            "Assunto: Decisão de 1ª instância proferida pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias.", 
            negrito=True
        )
        adicionar_paragrafo(doc, 
            f"Referência: Processo Administrativo Sancionador nº: {numero_processo} ", 
            negrito=True
        )
        doc.add_paragraph("\n")
        
        # Corpo do documento com conteúdo adaptado
        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, 
            f"Informamos que a Decisão em 1ª instância proferida pela Gerência-Geral de Portos, Aeroportos, Fronteiras e Recintos Alfandegados ou Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias, em {data_decisao.strftime('%d/%m/%Y')}, no processo administrativo sancionador em referência, foi revisada ou retratada no âmbito da Anvisa pelos motivos expostos abaixo."
        )
        doc.add_paragraph("\n")
        
        # Condições baseadas no motivo da revisão
        if motivo_revisao == "insuficiencia_provas":
            adicionar_paragrafo(doc, 
                "Foi constatado que não há comprovação suficiente nos autos do processo para afirmar que a recorrente cometeu a infração objeto da autuação em questão."
            )
        elif motivo_revisao == "prescricao":
            adicionar_paragrafo(doc, 
                f"Foi observado que da decisão condenatória recorrível proferida em {data_decisao.strftime('%d/%m/%Y')} até o ato seguinte capaz de interromper a prescrição (ex: notificação da decisão em {data_recebimento_notificacao.strftime('%d/%m/%Y')}) passaram-se mais de cinco anos sem que houvesse entre eles outro ato capaz de interromper o curso prescricional (documento que declarou a prescrição. Ex: NOTA n. 00014/2020/EI-M-ANVIS/ENAC/PGF/AGU)."
            )
        elif motivo_revisao == "extincao_empresa":
            if not data_extincao:
                raise ValueError("A data de extinção da empresa deve ser fornecida para o motivo 'extincao_empresa'.")
            adicionar_paragrafo(doc, 
                f"Foi constatado, ao longo dos procedimentos de cobrança administrativa, que a empresa em questão havia sido 'EXTINTA' na data de {data_extincao.strftime('%d/%m/%Y')}, conforme Certidão Simplificada e documento de Distrato Social fornecido pelo órgão de registro comercial - [Nome do Órgão]."
            )
        else:
            # Para outros motivos, conteúdo genérico
            adicionar_paragrafo(doc, 
                "Foi constatado que há razões adicionais para a revisão/retratação da decisão, conforme detalhado nos documentos anexos."
            )
        
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, 
            "Dessa forma, a decisão condenatória perdeu seus efeitos e o processo será arquivado."
        )
        doc.add_paragraph("\n")
        
        # Seção "Como Obter Cópia do Processo"
        adicionar_paragrafo(doc, "COMO OBTER CÓPIA DO PROCESSO?", negrito=True)
        adicionar_paragrafo(doc, 
            "Informações e pedidos de cópias devem ser solicitados exclusivamente pelos Canais de Atendimento da Anvisa (https://www.gov.br/anvisa/pt-br/canais_atendimento) ou pelo Serviço de Atendimento ao Cidadão (https://www.gov.br/anvisa/pt-br/acessoainformacao/sic)."
        )
        adicionar_paragrafo(doc, 
            "Os pedidos de cópia de processo devem informar o número do processo e a finalidade da cópia."
        )
        adicionar_paragrafo(doc, 
            "A cópia integral dos autos somente será concedida para o interessado direto no processo, ou seu representante devidamente constituído, cuja condição deve ser comprovada mediante a apresentação dos seguintes documentos:"
        )
        adicionar_paragrafo(doc, "1. Documento de identificação do autuado (se pessoa física) ou outorgado;")
        adicionar_paragrafo(doc, 
            "2. Procuração e documento de identificação do outorgado (advogado ou representante), caso seja ele o requerente. "
            "Somente serão aceitas procurações e substabelecimento assinados eletronicamente, com certificação digital no padrão da Infraestrutura de Chaves Públicas Brasileira (ICP-Brasil) ou pelo assinador Gov.br."
        )
        adicionar_paragrafo(doc, 
            "3. Contrato ou estatuto social da empresa, com a última alteração (se pessoa jurídica);"
        )
        adicionar_paragrafo(doc, 
            "4. Ata de eleição da atual diretoria quando a procuração estiver assinada por diretor que não conste como sócio da empresa (se pessoa jurídica);"
        )
        adicionar_paragrafo(doc, 
            "A ausência de quaisquer dos documentos acima ensejará o indeferimento sumário do pedido."
        )
        adicionar_paragrafo(doc, 
            "Terceiros não interessados diretamente no processo estão dispensados de apresentar documentação e terão acesso somente às cópias dos seguintes documentos: Auto de Infração, Manifestação da área autuante e Decisão."
        )
        adicionar_paragrafo(doc, f"\nInformações de contato: {email_selecionado}")

    except Exception as e:
        st.error(f"Erro ao gerar o documento no modelo 2: {e}")

def _gerar_modelo_3(doc, info, enderecos, numero_processo, usuario_nome, usuario_email, orgao_registro_comercial, email_selecionado):
    """
    Gera o Documento Word no Modelo 3.
    
    :param doc: Objeto Document do python-docx.
    :param info: Dicionário com informações extraídas.
    :param enderecos: Lista de dicionários com endereços.
    :param numero_processo: Número do processo formatado.
    :param usuario_nome: Nome do usuário que está gerando o documento.
    :param usuario_email: Email do usuário.
    :param orgao_registro_comercial: Órgão de registro comercial.
    :param email_selecionado: Email selecionado pelo usuário.
    """
    try:
        # Cabeçalho do documento
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        nome_autuado = info.get('nome_autuado', '[Nome não informado]')
        cnpj = info.get('cnpj', '')
        cpf = info.get('cpf', '')
        if cnpj:
            identificador = f"CNPJ: {cnpj}"
        elif cpf:
            identificador = f"CPF: {cpf}"
        else:
            identificador = "CNPJ/CPF: [Não informado]"
        adicionar_paragrafo(doc, f"{nome_autuado} – {identificador}")
        doc.add_paragraph("\n")

        # Informações de endereço
        for endereco in enderecos:
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Cidade: {endereco.get('cidade', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Bairro: {endereco.get('bairro', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Estado: {endereco.get('estado', '[Não informado]')}")
            adicionar_paragrafo(doc, f"CEP: {endereco.get('cep', '[Não informado]')}")
            doc.add_paragraph("\n")
        
        # Assunto e Referência em negrito
        adicionar_paragrafo(doc, "Assunto: Decisão proferida pela Diretoria Colegiada", negrito=True)
        adicionar_paragrafo(doc, f"Referência: Processo Administrativo Sancionador nº {numero_processo}", negrito=True)
        doc.add_paragraph("\n")
        
        # Corpo do documento com conteúdo específico
        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, 
            "Informamos que foi proferido julgamento da Diretoria Colegiada no processo administrativo sancionador em referência, conforme decisão em anexo, contra a qual não cabe recurso."
        )
        adicionar_paragrafo(doc, "\n")
        adicionar_paragrafo(doc, 
            "Em sendo mantida a penalidade de multa, esta notificação estará acompanhada de boleto bancário. Exceto para a decisão, cujo recurso tenha sido considerado intempestivo, um vez que o boleto será encaminhado pela Gerência de Gestão de Arrecadação – GEGAR."
        )
        adicionar_paragrafo(doc, "\n")
        adicionar_paragrafo(doc, 
            "O não pagamento do boleto, caso devido, acarretará, sucessivamente: i) a inscrição do devedor no Cadastro Informativo de Crédito não Quitado do Setor Público Federal (CADIN); ii) a inscrição do débito em dívida ativa da União; iii) o ajuizamento de ação de execução fiscal contra o devedor; e iv) a comunicação aos cartórios de registros de imóveis, dos devedores inscritos em dívida ativa ou execução fiscal."
        )
        adicionar_paragrafo(doc, "\n")
        adicionar_paragrafo(doc, 
            "Esclarecemos que, em caso de penalidade de multa, seu valor foi atualizado pela taxa Selic acumulada nos termos do art. 37-A da Lei 10.522/2002 e no art. 5º do Decreto-Lei 1.736/79."
        )
        adicionar_paragrafo(doc, "\n")
        
        # Seção "Informações e pedidos de cópias"
        adicionar_paragrafo(doc, 
            "Informações e pedidos de cópias podem ser solicitados pelos Canais de Atendimento da Anvisa (webchat, formulário eletrônico ou telefone 0800 642 9782), responsáveis por atender a esse tipo de demanda de forma centralizada. Os pedidos de cópia de PAS devem vir acompanhados dos documentos abaixo, sob pena de não serem atendidos:"
        )
        adicionar_paragrafo(doc, "\n")
        
        # Lista de documentos necessários
        documentos = [
            "Cópia autenticada da procuração/substabelecimento com firma reconhecida e poderes específicos para tal;",
            "Cópia do CPF e do RG do outorgado e do requerente, caso sejam pessoas distintas; e",
            "Cópia autenticada do contrato social/estatuto social, com a última alteração."
        ]
        
        for doc_item in documentos:
            adicionar_paragrafo(doc, f"- {doc_item}")
        adicionar_paragrafo(doc, "\n")
        
        adicionar_paragrafo(doc, 
            f"Por fim, esclarecemos que foi concedido aos autos por meio do Sistema Eletrônico de Informações (SEI), por 180 (cento e oitenta) dias, ao usuário: {usuario_nome} ({usuario_email})."
        )
        adicionar_paragrafo(doc, "\n")
        
        # Informações de contato
        adicionar_paragrafo(doc, f"\nInformações de contato: {email_selecionado}")
        
        # Encerramento do documento
        adicionar_paragrafo(doc, "Atenciosamente,")
        adicionar_paragrafo(doc, "\n")
        adicionar_paragrafo(doc, f"{usuario_nome}")
        
    except Exception as e:
        st.error(f"Erro ao gerar o documento no modelo 3: {e}")

###############################################################################
# Aplicação principal (Streamlit)
###############################################################################
def main():
    st.title("Gerador de Notificações SEI-Anvisa")

    # Seção de login
    st.sidebar.header("Informações de Login")
    if "username_input" not in st.session_state:
        st.session_state.username_input = ""
    if "password_input" not in st.session_state:
        st.session_state.password_input = ""

    # Campos de login
    st.session_state.username_input = st.sidebar.text_input("Usuário", value=st.session_state.username_input)
    st.session_state.password_input = st.sidebar.text_input("Senha", type="password", value=st.session_state.password_input)

    headless_option = st.sidebar.checkbox("Executar sem abrir o navegador (headless)?", value=True)
    
    # Seção de entrada do número do processo
    st.header("Processo Administrativo")
    if "process_number_input" not in st.session_state:
        st.session_state.process_number_input = ""

    st.session_state.process_number_input = st.text_input("Número do Processo", value=st.session_state.process_number_input)

    # Botão principal
    if st.button("Gerar Notificação e Extrair Dados"):
        if not st.session_state.username_input or not st.session_state.password_input or not st.session_state.process_number_input:
            st.error("Por favor, preencha todos os campos.")
        else:
            with st.spinner("Processando..."):
                try:
                    username_encrypted = cipher_suite.encrypt(st.session_state.username_input.encode('utf-8'))
                    password_encrypted = cipher_suite.encrypt(st.session_state.password_input.encode('utf-8'))

                    download_path = process_notification(
                        username_encrypted,
                        password_encrypted,
                        st.session_state.process_number_input,
                        headless=headless_option
                    )
                    st.success("PDF gerado/baixado com sucesso!")

                    if download_path:
                        pdf_file_name = os.path.basename(download_path)
                        numero_processo = extract_process_number(pdf_file_name)

                        text_final, enderecos_ocr = extract_text_with_best_ocr(download_path)

                        if text_final.strip():
                            st.success("Texto extraído com sucesso!")
                            info = extract_information_spacy(text_final)
                            addresses_ar_ais = extract_addresses_with_source(text_final)

                            # Unir endereços OCR e AR/AIS
                            all_addresses = addresses_ar_ais + enderecos_ocr

                            emails = extract_all_emails(info.get('emails', []))

                            # Guardar em session_state
                            st.session_state['info'] = info
                            st.session_state['addresses_raw'] = all_addresses
                            st.session_state['numero_processo'] = numero_processo
                            st.session_state['emails'] = emails

                except Exception as ex:
                    st.error(f"Ocorreu um erro: {ex}")

    # Só exibimos as informações extraídas se tivermos st.session_state populado
    if 'info' in st.session_state and 'addresses_raw' in st.session_state:
        st.subheader("Informações Extraídas")
        info = st.session_state['info']
        emails = st.session_state['emails']
        numero_processo = st.session_state['numero_processo']

        st.write(f"**Nome Autuado:** {info.get('nome_autuado', 'Não informado')}")
        if info.get('cnpj'):
            st.write(f"**CNPJ:** {info.get('cnpj')}")
        elif info.get('cpf'):
            st.write(f"**CPF:** {info.get('cpf')}")
        else:
            st.write("**CNPJ/CPF:** não encontrado ou inválido")
        st.write(f"**Emails:** {', '.join(emails) if emails else 'Não informado'}")
        st.write(f"**Sócios/Advogados:** {', '.join(info.get('socios_advogados', []))}")

        # Exibir endereços com a origem (não iremos mostrar no documento final, mas só para referência)
        st.subheader("Endereços Encontrados")
        if "addresses_edited" not in st.session_state:
            st.session_state['addresses_edited'] = st.session_state['addresses_raw'][:]

        # Permitir exclusão e edição de cada endereço
        for idx, end in enumerate(st.session_state['addresses_edited']):
            st.write(f"**Endereço {idx+1}**:")
            new_endereco = st.text_input(f"Endereço {idx+1}", value=end['endereco'], key=f"end_{idx}")
            new_cidade = st.text_input(f"Cidade {idx+1}", value=end['cidade'], key=f"cid_{idx}")
            new_bairro = st.text_input(f"Bairro {idx+1}", value=end['bairro'], key=f"bai_{idx}")
            new_estado = st.text_input(f"Estado {idx+1}", value=end['estado'], key=f"est_{idx}")
            new_cep = st.text_input(f"CEP {idx+1}", value=end['cep'], key=f"cep_{idx}")
            
            # Checkbox para excluir o endereço
            exclude_address = st.checkbox("Excluir este endereço?", key=f"excluir_{idx}", value=False)

            # Exibe a origem do endereço
            st.write(f"Origem do endereço: {end.get('source', 'Desconhecido')}")
            st.write("---")

            # Atualizar dicionário em session_state
            st.session_state['addresses_edited'][idx]['endereco'] = new_endereco
            st.session_state['addresses_edited'][idx]['cidade']   = new_cidade
            st.session_state['addresses_edited'][idx]['bairro']   = new_bairro
            st.session_state['addresses_edited'][idx]['estado']   = new_estado
            st.session_state['addresses_edited'][idx]['cep']      = new_cep
            st.session_state['addresses_edited'][idx]['excluded'] = exclude_address

        # Selecionar email
        st.subheader("Selecionar Email para Utilizar no Processo")
        if len(emails) > 0:
            if "selected_email" not in st.session_state:
                st.session_state.selected_email = emails[0]
            st.session_state.selected_email = st.selectbox("Selecione o email desejado:", emails, 
                                                          index=emails.index(st.session_state.selected_email) 
                                                          if st.session_state.selected_email in emails else 0)
        else:
            st.session_state.selected_email = "[Não informado]"
            st.write("Nenhum email encontrado.")

    # Segunda parte: geração de documentos
    if (
        'info' in st.session_state and 
        'addresses_edited' in st.session_state and 
        'numero_processo' in st.session_state
    ):
        st.subheader("Geração do Documento")
        modelo = st.selectbox(
            "Selecione o modelo desejado:",
            [
                "MODELO 1 - Notificação de decisões em 1ª instância",
                "MODELO 2 - Notificação de decisões revisadas/retratadas",
                "MODELO 3 - Notificação de decisão da DICOL"
            ]
        )

        if st.button("Gerar Documento Word"):
            try:
                doc = Document()
                info = st.session_state['info']

                # Antes de gerar, vamos filtrar os endereços que foram marcados como excluídos
                all_addresses = st.session_state['addresses_edited']
                final_addresses = [a for a in all_addresses if not a.get('excluded', False)]

                numero_processo = st.session_state['numero_processo']
                email_selecionado = st.session_state.get('selected_email', '[Não informado]')

                if "MODELO 1" in modelo:
                    _gerar_modelo_1(doc, info, final_addresses, numero_processo, email_selecionado)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    output_filename = f"Notificacao_{numero_processo}.docx"
                    st.download_button(
                        label="Baixar Documento",
                        data=buffer,
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("Documento (Modelo 1) gerado e pronto para download!")
                    return

                elif "MODELO 2" in modelo:
                    st.info("Preencha dados adicionais para o Modelo 2 (serão coletados agora):")
                    motivo_revisao = st.selectbox("Motivo da Revisão:", 
                        ["insuficiencia_provas", "prescricao", "extincao_empresa", "outros"], 
                        key="motivo_revisao_selectbox"
                    )
                    data_decisao = st.date_input("Data da Decisão:", key="data_decisao_input")
                    data_recebimento_notificacao = st.date_input("Data de Recebimento da Notificação:", key="data_receb_input")

                    data_extincao = None
                    if motivo_revisao == "extincao_empresa":
                        data_extincao = st.date_input("Data de Extinção da Empresa:", key="data_extincao_input")

                    if st.button("Gerar Modelo 2 Word"):
                        doc = Document()
                        _gerar_modelo_2(
                            doc,
                            info,
                            final_addresses,
                            numero_processo,
                            motivo_revisao,
                            data_decisao,
                            data_recebimento_notificacao,
                            data_extincao,
                            email_selecionado
                        )
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        output_filename = f"Notificacao_{numero_processo}_modelo2.docx"
                        st.download_button(
                            label="Baixar Documento",
                            data=buffer,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("Documento (Modelo 2) gerado e pronto para download!")
                    return

                elif "MODELO 3" in modelo:
                    st.info("Preencha dados adicionais para o Modelo 3 (serão coletados agora):")
                    usuario_nome = st.text_input("Nome do Usuário:", key="usuario_nome_input")
                    usuario_email = st.text_input("Email do Usuário:", key="usuario_email_input")
                    orgao_registro_comercial = st.text_input("Órgão de Registro Comercial:", key="orgao_registro_input")

                    if st.button("Gerar Modelo 3 Word"):
                        doc = Document()
                        _gerar_modelo_3(
                            doc,
                            info,
                            final_addresses,
                            numero_processo,
                            usuario_nome,
                            usuario_email,
                            orgao_registro_comercial,
                            email_selecionado
                        )
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        output_filename = f"Notificacao_{numero_processo}_modelo3.docx"
                        st.download_button(
                            label="Baixar Documento",
                            data=buffer,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("Documento (Modelo 3) gerado e pronto para download!")
                    return

            except Exception as ex:
                st.error(f"Ocorreu um erro ao gerar o documento: {ex}")

if __name__ == '__main__':
    try:
        nlp = spacy.load("pt_core_news_lg")
    except OSError:
        st.info("Modelo 'pt_core_news_lg' não encontrado. Instalando...")
        os.system("python -m spacy download pt_core_news_lg")
        nlp = spacy.load("pt_core_news_lg")

    main()
